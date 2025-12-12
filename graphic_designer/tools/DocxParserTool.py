from agency_swarm.tools import BaseTool
from pydantic import Field
from docx import Document
import os
import base64
from io import BytesIO

class DocxParserTool(BaseTool):
    """
    Extracts text, headings, paragraphs, and embedded images from DOCX files.
    Returns structured data with content hierarchy for creative brief generation.
    """
    file_path: str = Field(
        ..., description="Absolute path to the DOCX file to parse"
    )
    
    def run(self):
        """
        Parses DOCX file and extracts all text content and metadata.
        Returns a JSON string with structured content.
        """
        # Step 1: Validate file exists
        if not os.path.exists(self.file_path):
            return f"Error: File not found at {self.file_path}"
        
        # Step 2: Open and parse DOCX document
        try:
            doc = Document(self.file_path)
        except Exception as e:
            return f"Error parsing DOCX file: {str(e)}"
        
        # Step 3: Extract content
        results = {
            'headings': [],
            'paragraphs': [],
            'full_text': '',
            'images_count': 0,
            'metadata': {
                'filename': os.path.basename(self.file_path),
                'total_paragraphs': len(doc.paragraphs)
            }
        }
        
        all_text = []
        
        # Extract paragraphs and headings
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check if paragraph is a heading
            if para.style.name.startswith('Heading'):
                results['headings'].append({
                    'level': para.style.name,
                    'text': text
                })
            else:
                results['paragraphs'].append(text)
            
            all_text.append(text)
        
        results['full_text'] = '\n\n'.join(all_text)
        
        # Step 4: Count embedded images
        # Images in DOCX are in the document.part.rels
        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
        results['images_count'] = image_count
        
        # Step 5: Return structured results
        import json
        return json.dumps(results, indent=2)

if __name__ == "__main__":
    # Test case: Create and parse a sample DOCX
    import tempfile
    from docx import Document
    
    # Create test DOCX
    doc = Document()
    doc.add_heading('Sample Heading', level=1)
    doc.add_paragraph('This is a test paragraph for the DOCX parser.')
    doc.add_paragraph('Another paragraph with more content.')
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        doc.save(f.name)
        test_file = f.name
    
    tool = DocxParserTool(file_path=test_file)
    result = tool.run()
    print(result)
    
    # Cleanup
    os.unlink(test_file)
